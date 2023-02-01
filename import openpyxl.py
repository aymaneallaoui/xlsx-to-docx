import openpyxl
from docx import Document

# Load the Excel workbook
try:
    workbook = openpyxl.load_workbook("C:/Users/praxe/Documents/shitty program/january.xlsx")
except FileNotFoundError:
    print("The file example.xlsx could not be found.")
    exit(1)

# Select the active sheet
sheet = workbook.active

# Check if the sheet contains a table
if sheet.tables:
    # Get the first table in the sheet
    table_range = sheet.tables[0].ref

    # Create a new Word document
    document = Document()

    # Add a table to the Word document
    table = document.add_table(rows=table_range.height, cols=table_range.width)

    # Set the style for the table
    table.style = "Table Grid"

    # Loop through the rows and cells in the table range and add the values to the Word table
    for i, row in enumerate(sheet[table_range.min_row:table_range.max_row]):
        for j, cell in enumerate(row[table_range.min_col:table_range.max_col]):
            if cell.value:
                table.cell(i, j).text = str(cell.value)

    # Loop through the remaining cells in the sheet that are not part of the table
    for row in sheet[:table_range.min_row-1] + sheet[table_range.max_row+1:]:
        for cell in row:
            if cell.value:
                # Add the cell value as a new paragraph in the Word document
                document.add_paragraph(str(cell.value))

    # Save the Word document
    document.save("C:/Users/praxe/Documents/shitty program/january.docx")
else:
    print("The sheet does not contain a table.")
