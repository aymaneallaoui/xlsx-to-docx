import docx
import openpyxl


workbook = openpyxl.load_workbook("C:/Users/praxe/Documents/shitty program/january.xlsx")

lol = workbook.active

document = docx.Document()

for row in lol.iter_rows():
    paragraph = document.add_paragraph()
    for cell in row:
        if cell.value is not None:
            paragraph.add_run(str(cell.value)).bold = True

for table in lol.tables.values():
    paragraph.add_run(str(table.name)).bold = True

document.save('C:/Users/praxe/Documents/shitty program/january.docx')